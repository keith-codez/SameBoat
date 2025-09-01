# app/services/formatter.py
import json
import os
from datetime import datetime
from docx import Document

def save_output(title: str, content: str, tags: list[str]):
    slug = title.lower().replace(" ", "-")
    date_str = datetime.now().strftime("%Y-%m-%d")

    md_path = f"outputs/{date_str}-{slug}.md"
    json_path = f"outputs/{date_str}-{slug}.json"
    docx_path = f"outputs/{date_str}-{slug}.docx"

    # Save Markdown
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)

    # Save JSON
    data = {"title": title, "slug": slug, "content": content, "tags": tags}
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)

    # Save Word DOCX
    doc = Document()
    doc.add_heading(title, level=1)

    references = []
    in_references = False

    for line in content.split("\n"):
        stripped = line.strip()

        # Detect start of References section
        if stripped.lower().startswith("## references"):
            in_references = True
            continue

        if in_references:
            if stripped:  # collect non-empty lines as references
                references.append(stripped.lstrip("-* ").strip())
            continue

        # Convert headings
        if stripped.startswith("## "):
            doc.add_heading(stripped.replace("## ", ""), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped.replace("# ", ""), level=1)
        else:
            doc.add_paragraph(stripped)

    if tags:
        doc.add_heading("Tags", level=2)
        doc.add_paragraph(", ".join(tags))

    # Always include References section
    doc.add_heading("References", level=2)
    if references:
        for ref in references:
            doc.add_paragraph(ref, style="List Bullet")
    else:
        doc.add_paragraph("(No references provided)")

    doc.save(docx_path)

    return {"markdown": md_path, "json": json_path, "docx": docx_path}
