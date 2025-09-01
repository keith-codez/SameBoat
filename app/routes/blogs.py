from fastapi import APIRouter, HTTPException, Depends
from sqlalchemy.orm import Session
from fastapi.responses import FileResponse
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

from .. import crud, schemas
from ..database import get_db
from ..services.gpt_writer import generate_blog

router = APIRouter(prefix="/blogs", tags=["Blogs"])

# -------------------- CRUD --------------------
@router.get("/", response_model=list[schemas.BlogResponse])
def get_blogs(db: Session = Depends(get_db)):
    return crud.get_blogs(db)


@router.get("/{blog_id}", response_model=schemas.BlogResponse)
def read_blog(blog_id: int, db: Session = Depends(get_db)):
    db_blog = crud.get_blog(db, blog_id)
    if not db_blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    return db_blog
    

# -------------------- GENERATE --------------------
@router.post("/generate/{faq_id}")
def create_blog(faq_id: int, db: Session = Depends(get_db)):
    faq = crud.get_faq(db, faq_id)
    if not faq:
        raise HTTPException(status_code=404, detail="FAQ not found")

    blog_content = generate_blog(faq.question, faq.answer)

    output_dir = Path("outputs")
    output_dir.mkdir(exist_ok=True)
    slug = faq.question.lower().replace(" ", "-")

    doc = Document()
    doc.add_heading(faq.question, level=1)
    doc.add_paragraph(blog_content)

    file_path = output_dir / f"{slug}.docx"
    doc.save(file_path)

    return {"message": "Blog generated successfully", "file": str(file_path)}

# -------------------- EXPORT --------------------
EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

@router.get("/export/docx")
def export_blogs_docx(db: Session = Depends(get_db)):
    blogs = crud.get_blogs(db)
    file_path = os.path.join(EXPORT_DIR, "blogs_export.docx")

    doc = Document()
    doc.add_heading("Blogs", level=1)

    for blog in blogs:
        doc.add_heading(blog.title, level=2)
        doc.add_paragraph(blog.content)
        doc.add_paragraph("")

    doc.save(file_path)
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="blogs.docx"
    )

@router.get("/export/pdf")
def export_blogs_pdf(db: Session = Depends(get_db)):
    blogs = crud.get_blogs(db)
    file_path = os.path.join(EXPORT_DIR, "blogs_export.pdf")

    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    y = height - 50

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Blogs")
    y -= 30

    for blog in blogs:
        if y < 100:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 11)

        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, blog.title)
        y -= 20

        c.setFont("Helvetica", 11)
        for line in blog.content.splitlines():
            if y < 100:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 11)
            c.drawString(70, y, line)
            y -= 15
        y -= 20

    c.save()
    return FileResponse(file_path, media_type="application/pdf", filename="blogs.pdf")
