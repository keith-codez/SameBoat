from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from fastapi.responses import FileResponse
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

from .. import crud, schemas, models
from ..database import get_db

router = APIRouter(
    prefix="/faqs",
    tags=["FAQs"]
)

# -------------------- CRUD --------------------

# POST a new FAQ (AI disabled)
@router.post("/", response_model=schemas.FAQResponse)
def create_faq(faq: schemas.FAQCreate, db: Session = Depends(get_db)):
    db_faq = crud.create_faq(db, faq)
    db_faq.ai_response = None   # explicitly set to None
    db.commit()
    db.refresh(db_faq)
    return db_faq


@router.get("/", response_model=list[schemas.FAQResponse])
def get_faqs(db: Session = Depends(get_db)):
    return crud.get_faqs(db)


# GET FAQ by ID
@router.get("/{faq_id}", response_model=schemas.FAQResponse)
def read_faq(faq_id: int, db: Session = Depends(get_db)):
    db_faq = crud.get_faq(db, faq_id)
    if not db_faq:
        raise HTTPException(status_code=404, detail="FAQ not found")
    return db_faq


# -------------------- EXPORT --------------------

EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)


@router.get("/export/docx")
def export_faqs_docx(db: Session = Depends(get_db)):
    faqs = crud.get_faqs(db)
    file_path = os.path.join(EXPORT_DIR, "faqs_export.docx")

    doc = Document()
    doc.add_heading("FAQs", level=1)

    for faq in faqs:
        doc.add_heading(f"Q: {faq.question}", level=2)
        doc.add_paragraph(f"A: {faq.answer}")
        doc.add_paragraph("")

    doc.save(file_path)
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="faqs.docx"
    )


@router.get("/export/pdf")
def export_faqs_pdf(db: Session = Depends(get_db)):
    faqs = crud.get_faqs(db)
    file_path = os.path.join(EXPORT_DIR, "faqs_export.pdf")

    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    y = height - 50

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "FAQs")
    y -= 30

    c.setFont("Helvetica", 11)
    for faq in faqs:
        if y < 100:  # start new page if needed
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 11)

        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, f"Q: {faq.question}")
        y -= 15
        c.setFont("Helvetica", 11)
        for line in faq.answer.splitlines():
            c.drawString(70, y, f"A: {line}")
            y -= 15
        y -= 10

    c.save()
    return FileResponse(file_path, media_type="application/pdf", filename="faqs.pdf")
